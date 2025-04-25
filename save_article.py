import re
import fitz
import os
from docx import Document
from docx.shared import Inches
import logging
from collections import defaultdict
from datetime import datetime

logging.basicConfig(level=logging.INFO, 
                   format='%(asctime)s - %(levelname)s - %(message)s')

def get_todays_files():
    """Get today's newspaper files"""
    today = datetime.now().strftime("%d-%m-%Y")
    newspapers_dir = os.path.expanduser("~/Downloads/Newspapers")
    highlights_dir = os.path.expanduser("~/Downloads/Highlights/Text_highlights")
    
    files = []
    for pdf_file in os.listdir(newspapers_dir):
        if today in pdf_file and pdf_file.endswith('.pdf'):
            # Extract proper newspaper name
            if "Indian_Express" in pdf_file:
                newspaper_name = "Indian_Express"
            elif "The_Hindu" in pdf_file:
                newspaper_name = "The_Hindu"
            else:
                continue
                
            files.append({
                'name': newspaper_name,
                'pdf': os.path.join(newspapers_dir, pdf_file),
                'txt': os.path.join(highlights_dir, f"{newspaper_name}_{today}.txt")
            })
            
    logging.info(f"Found {len(files)} newspapers for today")
    return files

def find_correct_page(pdf, headline_text, given_page_num):
    # Split headline into words for flexible matching
    words = headline_text.split()
    # Take last 4 words for more accurate matching
    search_words = words[-4:] if len(words) >= 4 else words
    
    # Search range: check given page ± 3 pages to account for ads
    start_page = max(0, given_page_num - 3)
    end_page = min(len(pdf), given_page_num + 3)
    
    for page_num in range(start_page, end_page):
        page = pdf[page_num]
        page_text = page.get_text().lower()
        
        # Try exact match first
        if headline_text.lower() in page_text:
            return page_num
            
        # Try matching last 4 words
        if all(word.lower() in page_text for word in search_words):
            return page_num
            
    # Fallback to given page if no match found
    return given_page_num - 1

def extract_sections_from_txt(txt_file):
    page_sections = defaultdict(list)
    try:
        with open(txt_file, 'r', encoding='utf-8') as f:
            content = f.read()
            
        # Split content by "Page" and get sections
        parts = re.split(r': *Page \d+', content)
        page_numbers = re.findall(r': *Page (\d+)', content)
        
        # Process all parts including the first one
        for i in range(len(page_numbers)):
            text = parts[i].strip()
            page_num = int(page_numbers[i])
            # Skip if text is exactly either newspaper header
            if text and text not in ["The Indian EXPRESS", "THE HINDU"]:
                # Add headline to the page's list of headlines
                page_sections[page_num].append(text)
                logging.info(f"Found section for page {page_num}: {text}")
        
        return page_sections
    except Exception as e:
        logging.error(f"Error extracting sections: {str(e)}")
        raise


def save_to_docx(pdf_path, page_sections, output_docx):
    doc = Document()
    try:
        pdf = fitz.open(pdf_path)
        logging.info(f"Processing PDF with {len(pdf)} pages")

        # Process pages in order
        for page_num in sorted(page_sections.keys()):
            if page_num > len(pdf):
                logging.warning(f"Page {page_num} exceeds PDF length. Skipping.")
                continue

            logging.info(f"Processing page {page_num}")
            
            # Add section header
            doc.add_heading(f"Contents from Page {page_num}:", level=1)
            
            # Add all headlines for this page
            for headline in page_sections[page_num]:
                doc.add_paragraph(headline)
            
            # Find actual page containing any of the headlines
            actual_page = None
            for headline in page_sections[page_num]:
                found_page = find_correct_page(pdf, headline, page_num)
                if found_page is not None:
                    actual_page = found_page
                    break
            
            if actual_page is None:
                actual_page = page_num - 1
            
            # Add PDF page as image (only once per page)
            page = pdf[actual_page]
            pix = page.get_pixmap()
            temp_img = f'temp_page_{page_num}.png'
            pix.save(temp_img)
            
            doc.add_picture(temp_img, width=Inches(6))
            os.remove(temp_img)
            logging.info(f"Added content for page {page_num}")

        doc.save(output_docx)
        logging.info(f"Document saved successfully to {output_docx}")
        pdf.close()
    except Exception as e:
        logging.error(f"Error in PDF processing: {str(e)}")
        raise


def main():
    try:
        # Create To_read directory if it doesn't exist
        output_dir = os.path.expanduser('~/Downloads/To_read')
        os.makedirs(output_dir, exist_ok=True)
        
        # Get today's newspaper files
        todays_papers = get_todays_files()
        logging.info(f"Found {len(todays_papers)} newspapers for today")
        
        for paper in todays_papers:
            logging.info(f"Processing {paper['name']}")
            output_doc = os.path.join(output_dir, f"{paper['name']}_{datetime.now().strftime('%d-%m-%Y')}.docx")
            
            page_sections = extract_sections_from_txt(paper['txt'])
            save_to_docx(paper['pdf'], page_sections, output_doc)
            
        logging.info("All newspapers processed successfully")
    except Exception as e:
        logging.error(f"Main process error: {str(e)}")

if __name__ == "__main__":
    main()
