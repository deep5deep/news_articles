import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt

def get_base_path():
    github_workspace = os.getenv('GITHUB_WORKSPACE')
    if github_workspace:
        return os.path.join(github_workspace, 'Downloads')
    return os.path.join(os.path.expanduser('~'), 'Library', 'CloudStorage', 'OneDrive-IBM', 'Documents', 'Scripts', 'news', 'Downloads')

def create_to_read_folder():
    base_path = get_base_path()
    new_folder_path = os.path.join(base_path, 'To_read')
    os.makedirs(new_folder_path, exist_ok=True)
    return new_folder_path

def create_organized_text_folder():
    base_path = get_base_path()
    new_folder_path = os.path.join(base_path, 'Organized_text')
    os.makedirs(new_folder_path, exist_ok=True)
    return new_folder_path

def organize_text(input_file, output_docx_file, output_txt_file):
    with open(input_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # Remove all special characters and replace with spaces
    content = re.sub(r'[^a-zA-Z0-9\s]', ' ', content)
    content = ' '.join(content.split())  # Remove extra spaces

    # Identify the newspaper
    if "THE HINDU" in content:
        newspaper = "THE HINDU"
    elif "The Indian EXPRESS" in content:
        newspaper = "The Indian EXPRESS"
    else:
        print(f"Unknown newspaper format in {input_file}")
        return

    # Remove newspaper name for formatted output
    formatted_content = content.replace(newspaper, "").strip()
    
    # Divide formatted content into numbered points
    parts = re.split(r'(Page \d+)', formatted_content)
    points = [part.strip() for part in parts if part.strip()]

    # Organize content by pages
    page_content = {}

    for i in range(0, len(points), 2):
        if i + 1 < len(points) and points[i + 1].startswith("Page"):
            current_page = points[i + 1].split()[1]
            if current_page not in page_content:
                page_content[current_page] = []
            page_content[current_page].append(points[i])

    # Create a new Word document
    doc = Document()

    # Write organized content to the Word document and text file
    with open(output_txt_file, 'w', encoding='utf-8') as txt_file:
        for page, content in sorted(page_content.items(), key=lambda x: int(x[0])):
            doc.add_heading(f"Contents from Page {page}:", level=1)
            txt_file.write(f"Contents from Page {page}:\n")
            for item in content:
                doc.add_paragraph(item)
                txt_file.write(f"{item}\n")
            doc.add_paragraph()  # Add an empty paragraph for spacing
            txt_file.write("\n")

    # Save the Word document
    doc.save(output_docx_file)

def process_newspapers():
    text_highlights_folder = os.path.join(get_base_path(), 'Highlights', 'Text_highlights')
    to_read_folder = create_to_read_folder()
    organized_text_folder = create_organized_text_folder()

    today = datetime.now().strftime("%d-%m-%Y")
    newspapers = ['Indian_Express', 'The_Hindu']

    for newspaper in newspapers:
        input_file = f"{newspaper}_{today}.txt"
        output_docx_file = f"{newspaper}_{today}.docx"
        output_txt_file = f"{newspaper}_{today}.txt"

        input_path = os.path.join(text_highlights_folder, input_file)
        output_docx_path = os.path.join(to_read_folder, output_docx_file)
        output_txt_path = os.path.join(organized_text_folder, output_txt_file)

        if os.path.exists(input_path):
            organize_text(input_path, output_docx_path, output_txt_path)
            print(f"Organized text saved to {output_docx_path} and {output_txt_path}")
        else:
            print(f"Input file not found: {input_path}")

if __name__ == "__main__":
    process_newspapers()
