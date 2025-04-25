import os
import logging
from datetime import datetime
from docx import Document
from docx.shared import Inches
import fitz  # PyMuPDF

def get_base_path():
    if 'GITHUB_WORKSPACE' in os.environ:
        return os.path.join(os.environ['GITHUB_WORKSPACE'], 'Downloads')
    else:
        return os.path.join(os.path.expanduser('~'), 'Library', 'CloudStorage', 'OneDrive-IBM', 'Documents', 'Scripts', 'news', 'Downloads')

def find_correct_page(pdf, headline_text, given_page_num):
    """Finds the actual page in the PDF that contains the headline."""
    # Split headline into words for flexible matching
    words = headline_text.split()
    # Take last 4 words for more accurate matching
    search_words = words[-4:] if len(words) >= 4 else words
    
    # Search range: check given page Â± 3 pages to account for ads
    start_page = max(0, given_page_num - 3)
    end_page = min(len(pdf), given_page_num + 3)
    
    logging.info(f"Searching for '{headline_text}' around page {given_page_num} (pages {start_page} to {end_page})")
    
    for page_num in range(start_page, end_page):
        page = pdf[page_num]
        page_text = page.get_text().lower()
        
        # Log the current page being checked
        logging.info(f"Checking page {page_num + 1} for match...")
        
        # Try exact match first
        if headline_text.lower() in page_text:
            logging.info(f"Exact match found on page {page_num + 1}")
            return page_num
            
        # Try matching last 4 words
        if all(word.lower() in page_text for word in search_words):
            logging.info(f"Partial match found using last 4 words on page {page_num + 1}")
            return page_num
            
    # Fallback to given page if no match found
    logging.warning(f"No match found for '{headline_text}' in the range. Falling back to given page {given_page_num}.")
    return given_page_num - 1

def save_to_docx(pdf_path, page_sections, output_docx):
    doc = Document()
    try:
        pdf = fitz.open(pdf_path)
        logging.info(f"Processing PDF with {len(pdf)} pages")

        for page_num in sorted(page_sections.keys()):
            if page_num > len(pdf):
                logging.warning(f"Page {page_num} exceeds PDF length. Skipping.")
                continue

            logging.info(f"Processing page {page_num}")
            doc.add_heading(f"Contents from Page {page_num}:", level=1)

            for headline in page_sections[page_num]:
                doc.add_paragraph(headline)

            actual_page = find_correct_page(pdf, headline, page_num) or (page_num - 1)
            page = pdf[actual_page]
            pix = page.get_pixmap()
            temp_img = f'temp_page_{page_num}.png'
            pix.save(temp_img)
            
            doc.add_picture(temp_img, width=Inches(6))
            os.remove(temp_img)
            logging.info(f"Added content for page {page_num}")

        doc.save(output_docx)
        logging.info(f"Document saved successfully to {output_docx}")
        pdf.close()
    except Exception as e:
        logging.error(f"Error in PDF processing: {str(e)}")
        raise

def main():
    base_path = get_base_path()
    
    today = datetime.now().strftime("%d-%m-%Y")  # Get today's date in dd-mm-yyyy format
    newspapers = ['Indian_Express', 'The_Hindu']
    
    for newspaper in newspapers:
        organized_text_file = os.path.join(base_path, 'Organized_text', f'{newspaper}_{today}.txt')
        newspaper_pdf = os.path.join(base_path, 'Newspapers', f'{newspaper}_{today}.pdf')
        output_docx = os.path.join(base_path, 'To_read', f'{newspaper}_{today}.docx')
        
        # Ensure directories exist
        os.makedirs(os.path.dirname(organized_text_file), exist_ok=True)
        os.makedirs(os.path.dirname(newspaper_pdf), exist_ok=True)
        os.makedirs(os.path.dirname(output_docx), exist_ok=True)
        
        if os.path.exists(organized_text_file):
            if os.path.exists(newspaper_pdf):
                # Read organized text file and structure it by pages
                page_sections = {}
                with open(organized_text_file, 'r', encoding='utf-8') as f:
                    current_page = None
                    for line in f:
                        if line.strip():  # If line is not empty
                            if line.startswith("Contents from Page"):
                                current_page = int(line.split(":")[0].split()[-1])
                                page_sections[current_page] = []
                            elif current_page is not None:
                                page_sections[current_page].append(line.strip())

                save_to_docx(newspaper_pdf, page_sections, output_docx)
            else:
                logging.warning(f"PDF file not found for {newspaper}. Skipping PDF processing.")
                # Create a simple document with just the text content
                doc = Document()
                with open(organized_text_file, 'r', encoding='utf-8') as f:
                    doc.add_paragraph(f.read())
                doc.save(output_docx)
                logging.info(f"Created text-only document for {newspaper}")
        else:
            logging.warning(f"Organized text file not found for {newspaper}.")
        
        logging.info(f"Organized text file path: {organized_text_file}")
        logging.info(f"Newspaper PDF path: {newspaper_pdf}")

if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)
    main()